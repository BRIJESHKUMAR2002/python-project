import time
from flask import Flask, render_template, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
import markdown2
import subprocess
import os
import json
import ast
import openai
import openpyxl
from datetime import datetime
import threading
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus.doctemplate import PageBreak
from docx import Document
import re
import zipfile

app = Flask(__name__, static_url_path='/static/')
api_key = os.environ["OPENAI_API_KEY"] = "api_key"
Folder_path = "upload_file_/"

# SQLite database path
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///chatgpt_xlsx.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Create SQLite database connection
db = SQLAlchemy(app)
client = openai.OpenAI()


class Files(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255))
    uploaded_time = db.Column(db.DateTime, default=datetime.utcnow)
    status = db.Column(db.String(20))


class Prompt(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    text = db.Column(db.String(1000))


# Initialize SQLite database
with app.app_context():
    db.create_all()


def save_to_docx(cases, responses, file_path):
    pattern = re.compile(r'\*\*(.*?)\*\*' or "###")
    doc = Document()
    for i in range(len(cases)):
        case = cases[i]
        response = responses[i]
        paragraph = doc.add_paragraph()
        matches = re.finditer(pattern, response)
        # title_match = re.search(r'Title:', response)
        title_match = re.finditer('Title:', response)

        if title_match:
            run = paragraph.add_run(title_match)
            run.bold = True
        start_index = 0
        for match in matches:
            paragraph.add_run(response[start_index:match.start()])

            run = paragraph.add_run(match.group(1))
            run.bold = True

            start_index = match.end()

        paragraph.add_run(response[start_index:])
    doc.save(file_path)


def save_to_pdf(cases=None, responses=None, file_path=None):
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()

    # Define a custom style for bold text
    bold_style = styles['Normal'].clone('Bold')
    bold_style.fontName = 'Helvetica-Bold'

    # Adjust style to add space after paragraphs
    para_style = styles['Normal'].clone('ParaStyle')
    para_style.spaceAfter = 10  # Adjust the space after paragraphs

    content = []
    for i in range(len(cases)):
        case = cases[i]
        response = responses[i]

        # Write case information
        content.append(Paragraph(f'Case: {case}', bold_style))

        # Process the response
        paragraphs = response.split('\n\n')  # Split into paragraphs on double newline
        for paragraph in paragraphs:
            parts = []
            # Process each paragraph for bold and normal text
            for part in re.split(r'(\*\*.+?\*\*|###)', paragraph):
                if part.startswith('**') or part.startswith('###'):
                    # Remove markers and keep the text bold
                    part_text = part.replace('**', '').replace('###', '')
                    parts.append('<b>{}</b>'.format(part_text))
                else:
                    parts.append(part)
            # Join the parts with no space for inline bold, and create a paragraph
            paragraph_text = ''.join(parts)
            content.append(Paragraph(paragraph_text, para_style))

        # Add a page break after each case
        content.append(PageBreak())

    # Build the PDF document with all accumulated content
    doc.build(content)


def extract_data_from_file(file_path):
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active

    allCaseName_list = [sheet.cell(row=i, column=1).value for i in range(2, sheet.max_row + 1)]
    alldata_list = [sheet.cell(row=i, column=2).value for i in range(2, sheet.max_row + 1)]
    # print(len(allCaseName_list),len(alldata_list))
    # print(allCaseName_list[14],alldata_list[14])
    extracted_data = []
    for i in range(len(allCaseName_list)):
        case_name = allCaseName_list[i]
        case_data = alldata_list[i]
        if case_name is not None and case_data is not None:
            extracted_data.append({'case_name': [case_name], 'data': [case_data]})
        elif case_name is not None and case_data is None:
            extracted_data.append({'case_name': [case_name], 'data': [case_data]})
        else:
            pass
    return extracted_data


# USED THE CHATGPT FOR RESPONSE
def gpt_4_response(prompt):
    response = client.chat.completions.create(
        model="gpt-4-0125-preview",
        messages=[
            {"role": "system",
             "content": "You are a helpful intelligent assistant that generates the Response for the case."},
            {"role": "user", "content": prompt}
        ],
        temperature=1,
        max_tokens=4000,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    return response.choices[0].message.content


def empty_folder(folder_path):
    try:
        for file_name in os.listdir(folder_path):
            file_path = os.path.join(folder_path, file_name)
            if os.path.isfile(file_path):
                os.remove(file_path)
    except Exception as e:
        print(e)


@app.route('/')
def index():
    try:
        filedata_list = Files.query.all()
        prompt_list = Prompt.query.all()
        return render_template('index.html', filedata_list=filedata_list, prompt_list=prompt_list)
    except Exception as e:
        return f"An error occurred: {str(e)}"


def replace_text_in_paragraph(paragraph, placeholder, replacement):
    """
    Searches for a placeholder in a paragraph and replaces it with the given replacement text.
    Retains the font size, color, and style of the original text.
    Text between double asterisks will be converted to bold.
    """
    if placeholder in paragraph.text:
        # Iterate through each run in the paragraph
        for run in paragraph.runs:
            # Check if the placeholder is present in the run's text
            if placeholder in run.text:
                # Replace the placeholder with the replacement text
                replaced_text = run.text.replace(placeholder, replacement)
                # Clear the original run's text
                run.clear()
                # Split the replaced text based on double asterisks
                parts = replaced_text.split('**')
                # Initialize flag to track if we are within bolded section
                bold_section = False
                # Iterate through the parts and apply formatting as needed
                for part in parts:
                    # Add text to the paragraph
                    new_run = paragraph.add_run(part)
                    # Preserve original formatting attributes from the placeholder text
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.color.rgb = run.font.color.rgb
                    # Apply bold formatting if in bold section
                    if bold_section:
                        new_run.bold = True
                    # Toggle bold_section flag for the next part
                    bold_section = not bold_section


def fill_cv_template(template_path, output_path, replacements):
    """
    Fills a CV template with personal information.

    Args:
    - template_path: Path to the template document.
    - output_path: Path where the filled document will be saved.
    - replacements: Dictionary with placeholder text as keys and replacement text as values.
    """
    doc = Document(template_path)

    # Iterate through each paragraph in the document to replace placeholders
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            replace_text_in_paragraph(paragraph, placeholder, replacement)

    # If there are tables with placeholders, you can iterate through them as well
    # Example:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:

                    for placeholder, replacement in replacements.items():
                        replace_text_in_paragraph(paragraph, placeholder, replacement)

    doc.save(output_path)
    return output_path


########new


# def gpt_4_response(prompt):
#     client = OpenAI()
#
#     assistant = client.beta.assistants.create(
#         name="Content Writer",
#         instructions="You are a Expert Content writer. Write the a detailed content on the Case.",
#         model="gpt-4-0125-preview",
#     )
#     thread = client.beta.threads.create()
#
#     message = client.beta.threads.messages.create(
#         thread_id=thread.id,
#         role="user",
#         content=prompt
#     )
#
#     run = client.beta.threads.runs.create_and_poll(
#         thread_id=thread.id,
#         assistant_id=assistant.id,
#         instructions="Please address the user as Lawyer. The user want case details."
#     )
#
#     if run.status == 'completed':
#         messages = client.beta.threads.messages.list(
#             thread_id=thread.id
#         )
#         return messages
#     else:
#         print(run.status)
def zip_files(file_paths, output_zip_path):
    with zipfile.ZipFile(output_zip_path, 'w') as zipf:
        for file_path in file_paths:
            zipf.write(file_path, arcname=os.path.basename(file_path))
        print(f"Created zip archive: {output_zip_path}")


def main_fun(app, case_names, prompt_text, docx_file_path, uploaded_file=None, file_name_=None):
    print(case_names, "case_namescase_namescase_names")
    with app.app_context():
        try:
            ### FOR FILE---------------------------------------

            if isinstance(case_names, list):
                mains_file_for_pdf = []
                mains_file_for_docx = []

                for case_data in case_names:
                    case_name = case_data['case_name'][0]
                    print("case_name", case_name)
                    case_value = case_data['data'][0]
                    newresponse2 = None
                    if case_value is None or case_value == "":
                        prompt_text = f"""You are a content writer expert that write a details reports regarding the case data the are provide bellow Use the following instruction and provide details accordingly.

                            **create a summary/reports  for {case_name}:

                            Title of the summary should be the name of the CASES the year of the filing + {prompt_text} 

                            \n\n

                                                    """ + """
                            **** Note: Extract the following information form your knowledge and provide a dictionary format as given:

                   Case_data = {"Case": 'Case Title',
                    "year": 'Case Year',
                    "location": 'Location',
                    "Case No": 'Case Number',
                    "Judge Name": 'Judge Name',
                    "summary": 'summary in 100 words',
                    "Decision": 'Decision',
                    "Contracts": 'Contracts',
                    "Legal Significance": 'Legal Significance',
                    "Financial Judgment": 'Financial Judgment',
                    "key takeaways": 'key takeaways',
                    }


                    ****IF something is missing replace the value with Not given
                    **** Add the phrase "trade secret" in the complete content 5 times only.
                    **** Always bold 'trade secret'.                  
                    ****Remember do not add extra spaces between words or lines """f"""
                    ****Give the response of the given {case_name} only, Do not add previous data in your response if there is any Data.
"""

                    else:

                        prompt_text = f"""You are a content writer expert that write a details reports regarding the case data the are provide bellow Use the following instruction and provide details accordingly.
                            **create a summary/reports  for {case_name}:

                            Title of the summary should be the name of the CASES the year of the filing + {prompt_text} 

                            Also used the User Data as Reference If given, 

                            User Data:

                            {case_value}



                            """ + """\n\n\n
                            **** Note: Extract the following information form your knowledge and provide a dictionary format the keys and values in the string as given:

                    Case_data = {"Case": 'Case Title',
                    "year": 'Case Year',
                    "location": 'Location',
                    "Case No": 'Case Number',
                    "Judge Name": 'Judge Name',
                    "summary": 'summary in 100 words',
                    "Decision": 'Decision',
                    "Contracts": 'Contracts',
                    "Legal Significance": 'Legal Significance',
                    "Financial Judgment": 'Financial Judgment',
                    "key takeaways": 'key takeaways',
                    }



                    ****IF something is missing replace the value with Not given
                    **** Add the phrase "trade secret" in the complete content 5 times only.
                    **** Always bold 'trade secret'.
                    ****Remember do not add extra spaces between words or lines.""" + f"""
                    ****Give the response of the given {case_name} only, Do not add previous data in your response if there is any Data.
"""

                    print(case_name)
                    response = ''
                    response = gpt_4_response(prompt_text)
                    print("response :", response)
                    try:

                        # *****************************************************************************************************************************************************
                        print("response---", response, type(response))
                        responseList = response.split('{', 1)
                        # print(responseList)

                        responseList = responseList[1].rsplit('}', 1)
                        responsess = responseList[0]
                        responsess = responsess.replace("\n\'", '').replace("\n", '').replace('"', "'")
                        responsess = "{" + str(responsess) + "}"
                        # print(responsess, "777777777777")
                        responsell = json.dumps(responsess)
                        responsell = responsell.replace('\t\t\t', '')
                        responsell = responsell.replace(',}', '}')
                        responsell = responsell.replace("': '", """': " """)
                        responsell = responsell.replace("',", """ ", """)
                        responsell = responsell.replace("' ,", """ ", """)
                        responsell = responsell.replace("'}", """ "} """)
                        responsell = responsell.replace("['", """ [" """)
                        responsell = responsell.replace("']", """ "] """)
                        responsell = responsell.replace("' }", """ " } """)
                        responsell = responsell.replace(": *", """: " """)
                        responsell = responsell.replace('["', '''"''')
                        responsell = responsell.replace('"]', '''"''')
                        lastindex = responsell.rfind('"') or responsell.rfind("'")
                        newresponse2 = ast.literal_eval(responsell[1:lastindex])

                    except Exception as e:
                        print(e, "^^^^^^^^^^^^^^^^^")
                        pass

                        # Define the content to replace the placeholders
                    try:
                        replacements = {
                            'CaseName': str(newresponse2['Case']),
                            'SummaryDetails': str(newresponse2['summary']),
                            'DecisionDetails': str(newresponse2['Decision']),
                            'Legal_Significance_Details': str(newresponse2['Legal Significance']),
                            'Financial_Judgment_Details': str(newresponse2['Financial Judgment']),
                            'Takeways_Details': str(newresponse2['key takeaways']),
                            '{{year}}': str(newresponse2['year']),
                            'LOCATION': str(newresponse2['location']),
                            '{{judge}}': str(newresponse2['Judge Name']),
                            'CASENO.': str(newresponse2['Case No']),
                            # 'ContractsDetails': str(newresponse2['Contracts']),

                        }
                    except:
                        replacements = {
                            'CaseName': "Not given",
                            'SummaryDetails': response,
                            'DecisionDetails': "Not given",
                            'Legal_Significance_Details': "Not given",
                            'Financial_Judgment_Details': "Not given",
                            'Takeways_Details': "Not given",
                            '{{year}}': "Not given",
                            'LOCATION': "Not given",
                            '{{judge}}': "Not given",
                            'CASENO.': "Not given",
                            # 'ContractsDetails': "Not given",

                        }
                    template_path = 'new_template.docx'  # Update this path
                    output_path = case_name + ".docx"
                    if output_path.startswith(" "):
                        output_path = output_path[1:]

                    output_folder = 'static/'

                    if docx_file_path.endswith('.docx'):
                        # Generate DOC file
                        try:
                            fill_cv_template(template_path, output_folder + "docx/" + output_path, replacements)
                            file_name_ = case_name + ".docx"
                            print(os.getcwd() + "/static/docx/" + file_name_, "8888888888")

                            mains_file_for_docx.append(os.getcwd() + "/static/docx/" + file_name_)
                        except:
                            pass
                        if case_name == case_names[-1]['case_name'][0]:
                            zip_file_name = docx_file_path.replace(".docx", "Docx.zip")
                            zip_files(mains_file_for_docx, zip_file_name)
                        else:
                            pass

                    # Empty the upload folder
                    if docx_file_path.endswith('.pdf'):
                        try:
                            fill_cv_template(template_path, output_path, replacements)
                            convert_docx_to_pdf(output_path, output_folder + "pdf")
                            time.sleep(3)
                            print(output_folder, docx_file_path, '----------------------------------output_folder')
                            file_name_ = case_name + ".pdf"
                            mains_file_for_pdf.append(os.getcwd() + "/static/pdf/" + file_name_)

                            original_filename_without_extension = os.path.splitext(os.path.basename(output_path))[0]
                            default_pdf_path = os.path.join(output_folder, f"{original_filename_without_extension}.pdf")

                            # Rename or move the file to the desired output path
                            if os.path.exists(default_pdf_path):
                                os.rename(default_pdf_path, file_name_)
                                print(f"File successfully converted and saved as {file_name_}")
                            else:
                                pass
                            try:
                                current_dir = os.getcwd()
                                print(current_dir, "wwwwwwwwwwww")
                                full_path_to_file = os.path.join(current_dir, output_path)

                                if os.path.exists(full_path_to_file):
                                    os.remove(full_path_to_file)
                                    print(f"Successfully removed {full_path_to_file}")
                                else:
                                    print(f"The file {full_path_to_file} does not exist.")
                            except:
                                pass
                        except:
                            pass

                        if case_name == case_names[-1]['case_name'][0]:
                            zip_file_name = docx_file_path.replace(".pdf", "PDF.zip")
                            zip_files(mains_file_for_pdf, zip_file_name)
                        else:
                            pass
                    try:
                        empty_folder(Folder_path)
                        # Update the completed status for the processed file
                        print(zip_file_name, "PpPPPppp")
                        zip_file_ = zip_file_name.rsplit('/')[-1]
                        files_to_update = Files.query.filter_by(filename=zip_file_).all()
                        print(files_to_update, 'lllmm')
                        if files_to_update:
                            for file in files_to_update:
                                file.status = 'Completed'
                            db.session.commit()
                    except:
                        pass
            ### FOR NAME ONLY ------------------------------------------------
            else:

                prompt_text = f"""You are a content writer expert that write a details reports regarding the case data the are provide bellow Use the following instruction and provide details accordingly.
                            **create a summary/reports form your knowledge for CASES: {case_names}.

                    Title of the summary should be the name of the CASES the year of the filing. The detail instruction are following:\n""" + prompt_text + '''\n\n\n


                    **** Note: Extract the following information form your knowledge and provide a dictionary format as given:

                    Case_data = {"Case": Case Title,
                    "year": Case Year,
                    "location": Location,
                    "Case No": Case Number,
                    "Judge Name": Judge Name,
                    "summary": summary in 100 words,
                    "Decision": Decision,
                    "Contracts": Contracts,
                    "Legal Significance": Legal Significance,
                    "Financial Judgment": Financial Judgment,
                    "key takeaways": key takeaways,
                    }



                    ****IF something is missing replace the value with Not given
                    **** Add the phrase "trade secret" in the complete content 5 times only.
                    **** Always bold 'trade secret'.          
                    ****Use Markdown if needed to highlight text in summary or Decision or Contracts or Legal Significance or Financial Judgment or key takeaways.
                    ****Remember do not add extra spaces between words or lines """
                    '''

                newresponse2 = None
                prompt = prompt_text.replace('CASES', case_names)
                response = gpt_4_response(prompt)
                response = markdown2.markdown(response)
                try:
                    print("response---", response, type(response))
                    responseList = response.split('{', 1)
                    print(responseList)
                    responseList = responseList[1].rsplit('}', 1)
                    responsess = responseList[0]
                    responsess = responsess.replace("\n\'", '').replace("\n", '').replace('"', "'")
                    responsess = "{" + str(responsess) + "}"
                    print(responsess, "777777777777")
                    responsell = json.dumps(responsess)
                    responsell = responsell.replace('\t\t\t', '')
                    responsell = responsell.replace(',}', '}')
                    responsell = responsell.replace("': '", """': " """)
                    responsell = responsell.replace("',", """ ", """)
                    responsell = responsell.replace("' ,", """ ", """)
                    responsell = responsell.replace("'}", """ "} """)
                    responsell = responsell.replace("['", """ [" """)
                    responsell = responsell.replace("']", """ "] """)
                    responsell = responsell.replace("' }", """ " } """)
                    responsell = responsell.replace(": *", """: " """)

                    lastindex = responsell.rfind('"')
                    print(responsell[1:lastindex], type(responsell))

                    print(responsell.rfind('"'))
                    newresponse2 = ast.literal_eval(responsell[1:lastindex])
                    print(newresponse2.keys(),
                          '_____________--------------------______________------------______---newresponse2')

                except:
                    pass

                # Define the content to replace the placeholders
                try:
                    replacements = {
                        'CaseName': str(newresponse2['Case']),
                        'SummaryDetails': str(newresponse2['summary']),
                        'DecisionDetails': str(newresponse2['Decision']),
                        'Legal_Significance_Details': str(newresponse2['Legal Significance']),
                        'Financial_Judgment_Details': str(newresponse2['Financial Judgment']),
                        'Takeways_Details': str(newresponse2['key takeaways']),
                        '{{year}}': str(newresponse2['year']),
                        'LOCATION': str(newresponse2['location']),
                        '{{judge}}': str(newresponse2['Judge Name']),
                        'CASENO.': str(newresponse2['Case No']),
                        # 'ContractsDetails': str(newresponse2['Contracts']),

                    }
                except:
                    replacements = {
                        'CaseName': "Not given",
                        'SummaryDetails': response,
                        'DecisionDetails': "Not given",
                        'Legal_Significance_Details': "Not given",
                        'Financial_Judgment_Details': "Not given",
                        'Takeways_Details': "Not given",
                        '{{year}}': "Not given",
                        'LOCATION': "Not given",
                        '{{judge}}': "Not given",
                        'CASENO.': "Not given",
                        # 'ContractsDetails': "Not given",

                    }
                template_path = 'new_template.docx'  # Update this path
                output_path = case_names + ".docx"
                if output_path.startswith(" "):
                    output_path = output_path[1:]

                output_folder = 'static/'

                if docx_file_path.endswith('.pdf'):
                    fill_cv_template(template_path, output_path, replacements)
                    convert_docx_to_pdf(output_path, output_folder + "pdf")
                    time.sleep(6)
                    print(output_folder, docx_file_path, output_path, '----------------------------------output_folder')
                    file_name_ = case_names + ".pdf"
                    print(output_path, file_name_)

                    original_filename_without_extension = os.path.splitext(os.path.basename(output_path))[0]
                    default_pdf_path = os.path.join(output_folder, f"{original_filename_without_extension}.pdf")

                    # Rename or move the file to the desired output path
                    if os.path.exists(default_pdf_path):
                        os.rename(default_pdf_path, file_name_)
                        print(f"File successfully converted and saved as {file_name_}")
                    else:
                        pass

                    try:
                        current_dir = os.getcwd()
                        print(current_dir, "wwwwwwwwwwww")
                        full_path_to_file = os.path.join(current_dir, output_path)

                        if os.path.exists(full_path_to_file):
                            os.remove(full_path_to_file)
                            print(f"Successfully removed {full_path_to_file}")
                        else:
                            print(f"The file {full_path_to_file} does not exist.")
                    except:
                        pass

                else:
                    # Generate DOC file
                    fill_cv_template(template_path, output_folder + "docx/" + output_path, replacements)
                    file_name_ = case_names + ".docx"
                print(file_name_, output_folder, "000000000")

                files_to_update = Files.query.filter_by(filename=file_name_).all()
                print(files_to_update, 'kkk')
                if files_to_update:
                    for file in files_to_update:
                        file.status = 'Completed'
                    db.session.commit()

        except Exception as e:
            print(f"An error occurred: {str(e)}")


####mmmm
def convert_docx_to_pdf(docx_path, output_folder):
    try:
        # Make sure LibreOffice is installed and the `soffice` command is in your PATH
        subprocess.run(['soffice', '--convert-to', 'pdf', '--outdir', output_folder, docx_path], check=True)
        print("Conversion successful.")
    except subprocess.CalledProcessError as e:
        print("Error during conversion:", e)


#####mmmm
import os


@app.route('/process_message', methods=['POST'])
def process_message():
    try:
        file = request.files['training_file']
        print("-------------------------------------------------------------")
        # Check if the file is not None and has a filename
        if file and file.filename != '':
            uploaded_file = request.files['training_file']
            qa_file_path = os.path.join(Folder_path, uploaded_file.filename)
            uploaded_file.save(qa_file_path)

            file_type = request.form.get('file_type')
            print("---------------------------------------------")
            print("FILE TYPE :", file_type)
            if file_type == 'pdf':
                pdf_folder_path = 'static/pdf/'
                file_name_ = uploaded_file.filename.replace(".xlsx", 'PDF.zip')
            elif file_type == 'doc':
                docx_folder_path = 'static/docx/'

                file_name_ = uploaded_file.filename.replace(".xlsx", 'Docx.zip')
                # docx_file_path = os.path.join(docx_folder_path, file_name_)
            else:
                return 'Invalid file type'

            # Insert data into MySQL 'files' table
            try:
                file = Files(filename=file_name_, status='In Progress')
                db.session.add(file)
                db.session.commit()

            except Exception as e:
                print("An error occurred while inserting data:", e)

            print("===========================================================================")
            prompt_text = request.form.get('prompt')
            prompt = Prompt.query.first()
            if prompt:
                prompt.text = prompt_text
            else:
                prompt = Prompt(text=prompt_text)
                db.session.add(prompt)
            db.session.commit()

            docx_folder_path = 'static/docx/'

            file_type = request.form.get('file_type')
            print("FILE TYPE :", file_type)
            if file_type == 'pdf':
                pdf_folder_path = 'static/pdf/'
                file_name_ = uploaded_file.filename.replace("xlsx", 'pdf')
                docx_file_path = os.path.join(pdf_folder_path, file_name_)
            elif file_type == 'doc':
                file_name_ = uploaded_file.filename.replace("xlsx", 'docx')
                docx_file_path = os.path.join(docx_folder_path, file_name_)
            else:
                return 'Invalid file type'

            case_names = extract_data_from_file(qa_file_path)
            load_thread1 = threading.Thread(target=main_fun,
                                            args=(
                                                app, case_names, prompt_text, docx_file_path, uploaded_file,
                                                file_name_))
            load_thread1.start()

        else:
            case_name = request.form['client_name']
            prompt_text = request.form.get('prompt')
            prompt = Prompt.query.first()
            if prompt:
                prompt.text = prompt_text
            else:
                prompt = Prompt(text=prompt_text)
                db.session.add(prompt)
            db.session.commit()

            docx_folder_path = 'static/docx/'
            pdf_folder_path = 'static/pdf/'
            docx_file_path = ''
            file_type = request.form.get('file_type')
            print("FILE TYPE:", file_type)
            if file_type == 'pdf':
                file_name_ = case_name + '.pdf'
                print("File name", file_name_)
                docx_file_path = os.path.join(pdf_folder_path, file_name_)

                try:
                    file_name = Files(filename=file_name_, status='In Progress')
                    db.session.add(file_name)
                    db.session.commit()

                except Exception as e:
                    print("An error occurred while inserting data:", e)
            else:
                file_name_ = case_name + '.docx'
                docx_file_path = os.path.join(docx_folder_path, file_name_)
                print("File name", file_name_)

                try:
                    file_name = Files(filename=file_name_, status='In Progress')
                    db.session.add(file_name)
                    db.session.commit()
                except Exception as e:
                    print("An error occurred while inserting data:", e)

            case_names = case_name
            load_thread1 = threading.Thread(target=main_fun,
                                            args=(app, case_names, prompt_text, docx_file_path, file_name_))
            load_thread1.start()

        return render_template("Completed")

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download/<filename>')
def download(filename):
    print(filename, "download errrrrr")
    try:
        # filename = filename.replace("xlsx", 'docx') and filename.replace("xlsx", 'pdf')
        file_path = f'static/docx/{filename}'
        return send_file(file_path, as_attachment=True)
    except:
        file_path = f'static/pdf/{filename}'
        return send_file(file_path, as_attachment=True)


@app.route('/delete', methods=['POST'])
def delete_file():
    try:
        file_id = request.form.get('file_id')
        file = Files.query.filter_by(id=file_id).first()
        db.session.delete(file)
        db.session.commit()
        return "File deleted successfully"
    except Exception as e:
        return f"An error occurred: {str(e)}", 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=True)