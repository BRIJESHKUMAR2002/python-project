import time
from flask import Flask, render_template, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
import markdown2
import subprocess
import os
import json
import ast
import openai
import openpyxl
from datetime import datetime
import threading
import zipfile
from docx import Document
from docx.oxml.shared import OxmlElement, qn
import re
import csv

app = Flask(__name__, static_url_path='/static/')
api_key = os.environ["OPENAI_API_KEY"] = "api_key"
Folder_path = "upload_file_/"

# SQLite database path
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///chatgpt_xlsx.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Create SQLite database connection
db = SQLAlchemy(app)
client = openai.OpenAI()


class Files(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255))
    uploaded_time = db.Column(db.DateTime, default=datetime.utcnow)
    status = db.Column(db.String(20))


class Prompt(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    text = db.Column(db.String(1000))


# Initialize SQLite database
with app.app_context():
    db.create_all()


# Extract the data from file that user upload
def extract_data_from_file(file_path):
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active

    allCaseName_list = [sheet.cell(row=i, column=1).value for i in range(2, sheet.max_row + 1)]
    alldata_list = [sheet.cell(row=i, column=2).value for i in range(2, sheet.max_row + 1)]
    url_list = [sheet.cell(row=i, column=3).value for i in range(2, sheet.max_row + 1)]

    extracted_data = []
    for i in range(len(allCaseName_list)):
        case_name = allCaseName_list[i]
        case_data = alldata_list[i]
        url_data = url_list[i]
        if case_name is not None and case_data is not None and url_data is not None:
            extracted_data.append({'case_name': [case_name], 'data': [case_data], 'url_data': [url_data]})
        elif case_name is not None and case_data is None and url_data is None:
            extracted_data.append({'case_name': [case_name], 'data': [case_data], 'url_data': [url_data]})
        elif case_name is not None and case_data is not None and url_data is None:
            extracted_data.append({'case_name': [case_name], 'data': [case_data], 'url_data': [url_data]})
        elif case_name is not None and case_data is None and url_data is not None:
            extracted_data.append({'case_name': [case_name], 'data': [case_data], 'url_data': [url_data]})
        else:
            pass
    return extracted_data


# USED THE CHATGPT FOR RESPONSE
def gpt_4_response(prompt):
    response = client.chat.completions.create(
        model="gpt-4-0125-preview",
        messages=[
            {"role": "system",
             "content": """You are an expert Content writer. Follow the user prompt to generate the report for the 
                        content."""},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    return response.choices[0].message.content


# Empty the upload folder
def empty_folder(folder_path):
    try:
        for file_name in os.listdir(folder_path):
            file_path = os.path.join(folder_path, file_name)
            if os.path.isfile(file_path):
                os.remove(file_path)
    except Exception as e:
        print(e)


@app.route('/')
def index():
    try:
        filedata_list = Files.query.all()
        prompt_list = Prompt.query.all()
        return render_template('index.html', filedata_list=filedata_list, prompt_list=prompt_list)
    except Exception as e:
        return f"An error occurred: {str(e)}"


# Adding Hyperlink on the Case no.
def add_hyperlink(paragraph, text, url):
    """
    Adds a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                          is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run for the hyperlink and apply formatting
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# Replacing the text in the template
def replace_text_in_paragraph(paragraph, placeholder, replacement):
    """
    Searches for a placeholder in a paragraph and replaces it with the given replacement text or hyperlink.
    Handles bold text enclosed in **.
    """
    new_paragraph_text = ""
    for run in paragraph.runs:
        current_text = run.text
        if placeholder in current_text:
            # Replace the placeholder with replacement text
            current_text = current_text.replace(placeholder, replacement)
        # Handle bold formatting within **
        parts = re.split(r'(\*\*[^*]+\*\*)', current_text)  # Split by *text*
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Extract text without ** and set it as bold
                bold_text = part[2:-2]
                new_run = paragraph.add_run(bold_text)
                new_run.bold = True
            else:
                # Add non-bold text
                match2 = re.search(r'<a\s+href=["\'](.*?)["\']>(.*?)<\/a>', part)
                match = re.search(r'\[(.*?)\]\((.*?)\)', part)
                if match:
                    # If part matches hyperlink format
                    text, url = match.groups()
                    add_hyperlink(paragraph, text, url)

                elif match2:
                    url = match2.group(1)
                    text = match2.group(2)
                    add_hyperlink(paragraph, text, url)

                else:
                    new_run = paragraph.add_run(part)
                    # Preserve original formatting
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.color.rgb = run.font.color.rgb
        # Clear the original run after processing
        run.clear()


# Appending the Data into template
def fill_cv_template(template_path, output_path, replacements):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            replace_text_in_paragraph(paragraph, placeholder, replacement)

    # Handle tables similarly if needed
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in replacements.items():
                        replace_text_in_paragraph(paragraph, placeholder, replacement)

    doc.save(output_path)
    return output_path


# Creating Zip File for the files
def zip_files(file_paths, output_zip_path):
    with zipfile.ZipFile(output_zip_path, 'w') as zipf:
        for file_path in file_paths:
            zipf.write(file_path, arcname=os.path.basename(file_path))
        print(f"Created zip archive: {output_zip_path}")


def main_fun(app, case_names, user_prompt, docx_file_path, uploaded_file=None, file_name_=None, review_number=0):
    with app.app_context():
        try:
            # ----------------FOR FILE---------------------------------------

            if isinstance(case_names, list):
                mains_file_for_pdf = []
                mains_file_for_docx = []

                for case_data in case_names:
                    case_name = case_data['case_name'][0]
                    print("case_name", case_name)
                    case_value = case_data['data'][0]
                    url_value = case_data['url_data'][0]
                    newresponse2 = None
                    if case_value is None or case_value == "":
                        prompt_text = f""""You are a content writing expert tasked with producing detailed reports 
                        based on the case data provided below. Please adhere strictly to the instructions and use the 
                        user data as a reference to generate a comprehensive report.

                            Create a report {case_name}:
                            Title of the summary should be the name of the CASE.


                            **Instructions**:
                            - User Prompt: Follow the User Prompt guidelines closely to generate your response.

                            ##User Prompt:
                            {user_prompt}


                            Requirements:
                            - Do not omit any quantitative data, such as financial judgments.
                            - Provide the information extracted from your knowledge in a dictionary format with the specified keys and values.

                            **Expected Output Format**:
                            Case_data = {{
                                "Case": 'Case Title',
                                "year": 'Case Year',
                                "location": 'Location if available, otherwise "Not given"',
                                "Case No": 'Case Number if available, otherwise "Not given"',
                                "Judge Name": 'Judge Name if available, otherwise "Not given"',
                                "summary": 'Create a summary for the case.',
                                "Decision": 'Decision if available, otherwise "Not given"',
                                "Contracts": 'Information on contracts involved, otherwise "Not given"',
                                "Legal Significance": 'Legal significance of the case, otherwise "Not given"',
                                "Financial Judgment": 'Financial judgment amount, otherwise "Not given"',
                                "key takeaways": 'Key takeaways from the case, otherwise "Not given"',
                            }}

                            ##If something is missing, replace the value with 'Not given'.
                            ##Remember, do not add extra spaces between words or lines."""

                    else:
                        instruction = f'''You are a helpful intelligent assistant. Your task is to generate a comprehensive and detailed report in 2500 words of the provided case study.

                        Case Study: {case_value}

                        Cover all aspects of the case without omitting any significant information, especially any quantitative data.

                        Case Url: {url_value}

                        Ensure your report includes the following details in a paragraph:

                        Detailed Summary,Case,year,location,Case Number with url,Judge Name,Decision,Contracts,
                        Legal Significance,Financial Judgment,key takeaways Provide the information extracted from 
                        the case study in a Single paragraph format.'''

                        response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {"role": "system",
                                 "content": """You are an expert content writer. Follow the instructions below to 
                                            generate the report based on the provided content."""},
                                {"role": "user", "content": instruction}
                            ],
                            temperature=0.5,
                            top_p=1,
                            frequency_penalty=0,
                            presence_penalty=0
                        )
                        case_data_ = response.choices[0].message.content
                        print(f"--------------------{case_data_}---------------------------------")
                        prompt_text = f"""You are a report writing expert tasked with producing detailed reports 
                        based on the User Data provided below. Please adhere strictly to the instructions and use the 
                        User Data as a reference to generate a comprehensive report.
                            
                            Create a report for {case_name}:
                            
                            **Instructions**:
                            - **User Prompt**: Follow the User Prompt guidelines closely to generate your response.  
                                ##User Prompt:
                                {user_prompt}
                                                                    
                            - **User Data**: Utilize the data provided below as a reference for your response.
                                ##User Data:
                                {case_data_}
                                         
                            - **Case Url**: Use this url as a case url.
                                ##Case Url:
                                {url_value}  
                                                   
                            **Requirements**:
                            - Do not omit any quantitative data, such as financial judgments.
                            - Provide the information extracted from User Data in a dictionary format with the specified keys and values.
                            
                            **Expected Output Format in JSON **:
                            Case_data = {{
                                "Case": 'Case Title',
                                "year": 'Case Year',
                                "location": 'Location if available, otherwise "Not given"',
                                "Case No": '[Case Number](Case Url) if available, otherwise "Not given"',
                                "Judge Name": 'Judge Name if available, otherwise "Not given"',
                                "summary": 'Create a summary for the case.',
                                "Decision": 'Decision if available, otherwise "Not given"',
                                "Contracts": 'Information on contracts involved, otherwise "Not given"',
                                "Legal Significance": 'Legal significance of the case, otherwise "Not given"',
                                "Financial Judgment": 'Financial judgment amount, otherwise "Not given"',
                                "key takeaways": 'Key takeaways from the case, otherwise "Not given"',
                                "Case Url": Case url if available, otherwise "Not given",
                            }}
                            
                            ##If something is missing, replace the value with 'Not given'.
                            ##Remember, do not add extra spaces between words or lines.
                            """
                    response = gpt_4_response(prompt_text)

                    try:
                        print("response---", response, type(response))
                        responseList = response.split('{', 1)
                        responseList = responseList[1].rsplit('}', 1)
                        responsess = responseList[0]
                        responsess = responsess.replace("\n\'", '').replace("\n", '').replace('"', "'")
                        responsess = "{" + str(responsess) + "}"
                        responsell = json.dumps(responsess)
                        responsell = responsell.replace('\t\t\t', '')
                        responsell = responsell.replace(',}', '}')
                        responsell = responsell.replace("': '", """': " """)
                        responsell = responsell.replace("',", """ ", """)
                        responsell = responsell.replace("' ,", """ ", """)
                        responsell = responsell.replace("'}", """ "} """)
                        responsell = responsell.replace("['", """ [" """)
                        responsell = responsell.replace("']", """ "] """)
                        responsell = responsell.replace("' }", """ " } """)
                        responsell = responsell.replace(": *", """: " """)
                        responsell = responsell.replace('["', '''"''')
                        responsell = responsell.replace('"]', '''"''')
                        lastindex = responsell.rfind('"') or responsell.rfind("'")
                        newresponse2 = ast.literal_eval(responsell[1:lastindex])
                    except Exception as e:
                        print(e, "^^^^^^^^^^^^^^^^^")
                        pass
                    try:
                        # Define the content to replace the placeholders
                        replacements = {
                            'review': str(review_number),
                            'CaseName': str(newresponse2['Case']),
                            'SummaryDetails': str(newresponse2['summary']),
                            'DecisionDetails': str(newresponse2['Decision']),
                            'Legal_Significance_Details': str(newresponse2['Legal Significance']),
                            'Financial_Judgment_Details': str(newresponse2['Financial Judgment']),
                            'Takeways_Details': str(newresponse2['key takeaways']),
                            '{{year}}': str(newresponse2['year']),
                            'LOCATION': str(newresponse2['location']),
                            '{{judge}}': str(newresponse2['Judge Name']),
                            'CASENO.': str(newresponse2['Case No']),
                        }
                        csv_data = {
                            'CASE LAW REVIEW #': str(review_number),
                            'Title': str(newresponse2['Case']),
                            'Year': str(newresponse2['year']),
                            'Place': str(newresponse2['location']),
                            'Judge': str(newresponse2['Judge Name']),
                            'Case': str(newresponse2['Case No']),
                            'Link of the case': str(newresponse2['Case Url']),
                            'Summary': str(newresponse2['summary']),
                            'Decision': str(newresponse2['Decision']),
                            'Legal Significance': str(newresponse2['Legal Significance']),
                            'Financial Judgment': str(newresponse2['Financial Judgment']),
                            'Takeways Details': str(newresponse2['key takeaways']),
                        }
                    except Exception as e:
                        replacements = {
                            'review': str(review_number),
                            'Title': "Not given",
                            'SummaryDetails': response,
                            'DecisionDetails': "Not given",
                            'Legal_Significance_Details': "Not given",
                            'Financial_Judgment_Details': "Not given",
                            'Takeways_Details': "Not given",
                            '{{year}}': "Not given",
                            'LOCATION': "Not given",
                            '{{judge}}': "Not given",
                            'CASENO.': "Not given",
                        }
                        csv_data = {
                            'CASE LAW REVIEW #': str(review_number),
                            'Title': "Not given",
                            'Year': "Not given",
                            'Place': "Not given",
                            'Judge': "Not given",
                            'Case': "Not given",
                            'Link of the case': "Not given",
                            'Summary': str(response),
                            'Decision': "Not given",
                            'Legal Significance': "Not given",
                            'Financial Judgment': "Not given",
                            'Takeways Details': "Not given",
                        }
                    template_path = 'template.docx'  # Update this path
                    if len(case_name) > 100:
                        desired_part = case_name.split(";")[0].strip()
                        output_path = desired_part + ".docx"
                    else:
                        output_path = case_name + ".docx"

                    if output_path.startswith(" "):
                        output_path = output_path[1:]

                    output_folder = 'static/'

                    if docx_file_path.endswith('.csv'):
                        def append_to_csv(file_path, data):
                            # Check if the file exists
                            file_exists = os.path.isfile(file_path)

                            # Open the file in append mode ('a'), create it if it doesn't exist
                            with open(file_path, 'a', newline='', encoding='utf-8') as csvfile:
                                fieldnames = list(data.keys())
                                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

                                # Write the header only if the file didn't exist before
                                if not file_exists:
                                    writer.writeheader()

                                # Write the data
                                writer.writerow(data)

                        append_to_csv(docx_file_path, csv_data)
                    else:
                        if docx_file_path.endswith('.docx'):
                            # Generate DOC file
                            try:
                                fill_cv_template(template_path, output_folder + "docx/" + output_path, replacements)
                                if len(case_name) > 100:
                                    desired_part = case_name.split(";")[0].strip()
                                    file_name_ = desired_part + ".docx"
                                else:
                                    file_name_ = case_name + ".docx"
                                print(os.getcwd() + "/static/docx/" + file_name_, "8888888888")
                                mains_file_for_docx.append(os.getcwd() + "/static/docx/" + file_name_)
                            except:
                                pass
                            if case_name == case_names[-1]['case_name'][0]:
                                zip_file_name = docx_file_path.replace(".docx", "Docx.zip")
                                zip_files(mains_file_for_docx, zip_file_name)
                            else:
                                pass

                        # Empty the upload folder
                        if docx_file_path.endswith('.pdf'):
                            try:
                                fill_cv_template(template_path, output_path, replacements)
                                convert_docx_to_pdf(output_path, output_folder + "pdf")
                                time.sleep(3)
                                print(output_folder, docx_file_path, '----------------------------------output_folder')
                                if len(case_name) > 150:
                                    desired_part = case_name.split(";")[0].strip()
                                    file_name_ = desired_part + ".pdf"
                                else:
                                    file_name_ = case_name + ".pdf"
                                mains_file_for_pdf.append(os.getcwd() + "/static/pdf/" + file_name_)
                                original_filename_without_extension = os.path.splitext(os.path.basename(output_path))[0]
                                default_pdf_path = os.path.join(output_folder,
                                                                f"{original_filename_without_extension}.pdf")

                                # Rename or move the file to the desired output path
                                if os.path.exists(default_pdf_path):
                                    os.rename(default_pdf_path, file_name_)
                                    print(f"File successfully converted and saved as {file_name_}")
                                else:
                                    pass
                                try:
                                    current_dir = os.getcwd()
                                    print(current_dir, "wwwwwwwwwwww")
                                    full_path_to_file = os.path.join(current_dir, output_path)

                                    if os.path.exists(full_path_to_file):
                                        os.remove(full_path_to_file)
                                        print(f"Successfully removed {full_path_to_file}")
                                    else:
                                        print(f"The file {full_path_to_file} does not exist.")
                                except Exception as e:
                                    print(e, 'current_dircurrent_dir')
                                    pass
                            except Exception as e:
                                print(e, 'endswith.pdf')
                                pass

                            if case_name == case_names[-1]['case_name'][0]:
                                zip_file_name = docx_file_path.replace(".pdf", "PDF.zip")
                                print(zip_file_name, 'zip_file_namezip_file_namezip_file_name')
                                print(mains_file_for_pdf, 'mains_file_for_pdfmains_file_for_pdf')
                                zip_files(mains_file_for_pdf, zip_file_name)
                            else:
                                pass
                    try:
                        empty_folder(Folder_path)
                        # Update the completed status for the processed file
                        print(zip_file_name, "PpPPPppp")
                        zip_file_ = zip_file_name.rsplit('/')[-1]
                        files_to_update = Files.query.filter_by(filename=zip_file_).all()
                        print(files_to_update, 'lllmm')
                        if files_to_update:
                            for file in files_to_update:
                                file.status = 'Completed'
                            db.session.commit()
                    except:
                        empty_folder(Folder_path)
                        # Update the completed status for the processed file
                        print(docx_file_path, 'dfdhfsghfdfsghdfsghdfsghdfsghdfsgh')
                        csv_pathhh = docx_file_path.split('/')[-1]
                        print(csv_pathhh, 'dfdhfsghfdfsghdfsghdfsghdfsghdfsgh')
                        files_to_update = Files.query.filter_by(filename=csv_pathhh).all()
                        print(files_to_update, 'lllmm')
                        if files_to_update:
                            for file in files_to_update:
                                file.status = 'Completed'
                            db.session.commit()
                    jj = int(review_number) + 1
                    ll = len(str(jj))
                    dd = len(str(review_number))
                    if dd != ll:
                        # Desired length of the output string (including leading zeros)
                        desired_length = dd
                        # Add leading zeros using string formatting
                        output_string = "{:0>{}}".format(str(jj), desired_length)
                        review_number = output_string
                    else:
                        review_number = jj

            # --------------------FOR NAME ONLY ----------------------------------
            else:
                prompt_text = f"""You are a content writer expert that write a details reports regarding the case 
                data the are provide bellow Use the following instruction and provide details accordingly. **create a 
                summary/reports form your knowledge for CASES: {case_names}.

                Title of the summary should be the name of the CASES the year of the filing. The detail instruction 
                are following:\n""" + user_prompt + """\n\n\n


                **** Note: Extract the following information form your knowledge and provide a dictionary format as given:

                Case_data = {"Case": Case Title,
                "year": Case Year,
                "location": Location,
                "Case No": Case Number,
                "Judge Name": Judge Name,
                "summary": summary in 100 words,
                "Decision": Decision,
                "Contracts": Contracts,
                "Legal Significance": Legal Significance,
                "Financial Judgment": Financial Judgment,
                "key takeaways": key takeaways,
                }

                ****IF something is missing replace the value with Not given
                **** Add the phrase "trade secret" in the complete content 5 times only.
                **** Always bold 'trade secret'.          
                ****Use Markdown if needed to highlight text in summary or Decision or Contracts or Legal Significance or Financial Judgment or key takeaways.
                ****Remember do not add extra spaces between words or lines.
                """

                newresponse2 = None
                prompt = prompt_text.replace('CASES', case_names)
                response = markdown2.markdown(gpt_4_response(prompt))
                try:
                    # print("response---", response, type(response))
                    responseList = response.split('{', 1)
                    # print(responseList)
                    responseList = responseList[1].rsplit('}', 1)
                    responsess = responseList[0]
                    responsess = responsess.replace("\n\'", '').replace("\n", '').replace('"', "'")
                    responsess = "{" + str(responsess) + "}"
                    print(responsess, "777777777777")
                    responsell = json.dumps(responsess)
                    responsell = responsell.replace('\t\t\t', '')
                    responsell = responsell.replace(',}', '}')
                    responsell = responsell.replace("': '", """': " """)
                    responsell = responsell.replace("',", """ ", """)
                    responsell = responsell.replace("' ,", """ ", """)
                    responsell = responsell.replace("'}", """ "} """)
                    responsell = responsell.replace("['", """ [" """)
                    responsell = responsell.replace("']", """ "] """)
                    responsell = responsell.replace("' }", """ " } """)
                    responsell = responsell.replace(": *", """: " """)

                    lastindex = responsell.rfind('"')
                    print(responsell[1:lastindex], type(responsell))

                    print(responsell.rfind('"'))
                    newresponse2 = ast.literal_eval(responsell[1:lastindex])
                    print(newresponse2.keys(),
                          '_____________--------------------______________------------______---newresponse2')

                except:
                    pass

                # Define the content to replace the placeholders

                try:
                    replacements = {
                        'review': str(review_number),
                        'CaseName': str(newresponse2['Case']),
                        'SummaryDetails': str(newresponse2['summary']),
                        'DecisionDetails': str(newresponse2['Decision']),
                        'Legal_Significance_Details': str(newresponse2['Legal Significance']),
                        'Financial_Judgment_Details': str(newresponse2['Financial Judgment']),
                        'Takeways_Details': str(newresponse2['key takeaways']),
                        '{{year}}': str(newresponse2['year']),
                        'LOCATION': str(newresponse2['location']),
                        '{{judge}}': str(newresponse2['Judge Name']),
                        'CASENO.': str(newresponse2['Case No']),
                        # 'ContractsDetails': str(newresponse2['Contracts']),

                    }
                    csv_data = {
                        'CASE LAW REVIEW #': str(review_number),
                        'Title': str(newresponse2['Case']),
                        'Year': str(newresponse2['year']),
                        'Place': str(newresponse2['location']),
                        'Judge': str(newresponse2['Judge Name']),
                        'Case': str(newresponse2['Case No']),
                        'Summary': str(newresponse2['summary']),
                        'Decision': str(newresponse2['Decision']),
                        'Legal Significance': str(newresponse2['Legal Significance']),
                        'Financial Judgment': str(newresponse2['Financial Judgment']),
                        'Takeways Details': str(newresponse2['key takeaways']),
                    }
                except:
                    replacements = {
                        'review': str(review_number),
                        'CaseName': "Not given",
                        'SummaryDetails': response,
                        'DecisionDetails': "Not given",
                        'Legal_Significance_Details': "Not given",
                        'Financial_Judgment_Details': "Not given",
                        'Takeways_Details': "Not given",
                        '{{year}}': "Not given",
                        'LOCATION': "Not given",
                        '{{judge}}': "Not given",
                        'CASENO.': "Not given",
                        # 'ContractsDetails': "Not given",

                    }
                    csv_data = {
                        'CASE LAW REVIEW #': str(review_number),
                        'Title': "Not given",
                        'Year': "Not given",
                        'Place': "Not given",
                        'Judge': "Not given",
                        'Case': "Not given",
                        'Summary': response,
                        'Decision': "Not given",
                        'Legal Significance': "Not given",
                        'Financial Judgment': "Not given",
                        'Takeways Details': "Not given",
                    }

                template_path = 'template.docx'  # Update this path
                output_path = case_names + ".docx"
                if output_path.startswith(" "):
                    output_path = output_path[1:]

                output_folder = 'static/'
                if docx_file_path.endswith('.csv'):
                    def append_to_csv(file_path, data):
                        # Check if the file exists
                        file_exists = os.path.isfile(file_path)

                        # Open the file in append mode ('a'), create it if it doesn't exist
                        with open(file_path, 'a', newline='', encoding='utf-8') as csvfile:
                            fieldnames = list(data.keys())
                            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

                            # Write the header only if the file didn't exist before
                            if not file_exists:
                                writer.writeheader()

                            # Write the data
                            writer.writerow(data)
                    print("CSV DATA::::::::::::::::::::::::::::::::::::::::::::::::::")
                    append_to_csv(docx_file_path, csv_data)
                    file_name_ = case_names + ".csv"
                else:
                    if docx_file_path.endswith('.pdf'):
                        fill_cv_template(template_path, output_path, replacements)
                        convert_docx_to_pdf(output_path, output_folder + "pdf")
                        time.sleep(6)
                        print(output_folder, docx_file_path, output_path,
                              '----------------------------------output_folder')
                        file_name_ = case_names + ".pdf"
                        print(output_path, file_name_)

                        original_filename_without_extension = os.path.splitext(os.path.basename(output_path))[0]
                        default_pdf_path = os.path.join(output_folder, f"{original_filename_without_extension}.pdf")

                        # Rename or move the file to the desired output path
                        if os.path.exists(default_pdf_path):
                            os.rename(default_pdf_path, file_name_)
                            print(f"File successfully converted and saved as {file_name_}")
                        else:
                            pass

                        try:
                            current_dir = os.getcwd()
                            print(current_dir, "wwwwwwwwwwww")
                            full_path_to_file = os.path.join(current_dir, output_path)

                            if os.path.exists(full_path_to_file):
                                os.remove(full_path_to_file)
                                print(f"Successfully removed {full_path_to_file}")
                            else:
                                print(f"The file {full_path_to_file} does not exist.")
                        except:
                            pass

                    else:
                        # Generate DOC file
                        fill_cv_template(template_path, output_folder + "docx/" + output_path, replacements)
                        file_name_ = case_names + ".docx"
                print(file_name_, output_folder, "000000000")
                files_to_update = Files.query.filter_by(filename=file_name_).all()

                print(files_to_update, 'kkk')
                if files_to_update:
                    for file in files_to_update:
                        file.status = 'Completed'
                    db.session.commit()

        except Exception as e:
            print(f"An error occurred: {str(e)}")


# Converting docx to pdf
def convert_docx_to_pdf(docx_path, output_folder):
    try:
        # Make sure LibreOffice is installed and the `soffice` command is in your PATH
        subprocess.run(['soffice', '--convert-to', 'pdf', '--outdir', output_folder, docx_path], check=True)
        print("Conversion successful.")
    except subprocess.CalledProcessError as e:
        print("Error during conversion:", e)


@app.route('/process_message', methods=['POST'])
def process_message():
    try:
        file = request.files['training_file']
        review_number = request.form.get('review')
        #### This is for xlsx file
        # Check if the file is not None and has a filename
        if file and file.filename != '':
            uploaded_file = request.files['training_file']
            qa_file_path = os.path.join(Folder_path, uploaded_file.filename)
            uploaded_file.save(qa_file_path)

            file_type = request.form.get('file_type')
            print("FILE TYPE :", file_type)
            if file_type == 'pdf':
                file_name_ = uploaded_file.filename.replace(".xlsx", 'PDF.zip')
            elif file_type == 'doc':
                file_name_ = uploaded_file.filename.replace(".xlsx", 'Docx.zip')
            elif file_type == 'csv':
                file_name_ = uploaded_file.filename.replace(".xlsx", '.csv')
            else:
                return 'Invalid file type'

            # Insert data into MySQL 'files' table
            try:
                file = Files(filename=file_name_, status='In Progress')
                db.session.add(file)
                db.session.commit()

            except Exception as e:
                print("An error occurred while inserting data:", e)

            prompt_text = request.form.get('prompt')
            prompt = Prompt.query.first()
            if prompt:
                prompt.text = prompt_text
            else:
                prompt = Prompt(text=prompt_text)
                db.session.add(prompt)
            db.session.commit()

            docx_folder_path = 'static/docx/'
            csv_folder_path = 'static/csv/'
            file_type = request.form.get('file_type')
            print("FILE TYPE :", file_type)
            if file_type == 'pdf':
                pdf_folder_path = 'static/pdf/'
                file_name_ = uploaded_file.filename.replace("xlsx", 'pdf')
                docx_file_path = os.path.join(pdf_folder_path, file_name_)
                case_names = extract_data_from_file(qa_file_path)
                load_thread1 = threading.Thread(target=main_fun,
                                                args=(
                                                    app, case_names, prompt_text, docx_file_path, uploaded_file,
                                                    file_name_,
                                                    review_number))
                load_thread1.start()
            elif file_type == 'doc':
                file_name_ = uploaded_file.filename.replace("xlsx", 'docx')
                docx_file_path = os.path.join(docx_folder_path, file_name_)
                case_names = extract_data_from_file(qa_file_path)
                load_thread1 = threading.Thread(target=main_fun,
                                                args=(
                                                    app, case_names, prompt_text, docx_file_path, uploaded_file,
                                                    file_name_,
                                                    review_number))
                load_thread1.start()
            elif file_type == 'csv':
                file_name_ = uploaded_file.filename.replace("xlsx", 'csv')
                docx_file_path = os.path.join(csv_folder_path, file_name_)
                case_names = extract_data_from_file(qa_file_path)
                load_thread1 = threading.Thread(target=main_fun,
                                                args=(
                                                    app, case_names, prompt_text, docx_file_path, uploaded_file,
                                                    file_name_,
                                                    review_number))
                load_thread1.start()
            else:
                return 'Invalid file type'



        ###This is for only case name
        else:
            case_name = request.form['client_name']
            prompt_text = request.form.get('prompt')
            prompt = Prompt.query.first()
            if prompt:
                prompt.text = prompt_text
            else:
                prompt = Prompt(text=prompt_text)
                db.session.add(prompt)
            db.session.commit()

            docx_folder_path = 'static/docx/'
            csv_folder_path = 'static/csv/'
            pdf_folder_path = 'static/pdf/'
            docx_file_path = ''
            file_type = request.form.get('file_type')
            print("FILE TYPE:", file_type)
            if file_type == 'pdf':
                file_name_ = case_name + '.pdf'
                print("File name", file_name_)
                docx_file_path = os.path.join(pdf_folder_path, file_name_)

                try:
                    file_name = Files(filename=file_name_, status='In Progress')
                    db.session.add(file_name)
                    db.session.commit()

                except Exception as e:
                    print("An error occurred while inserting data:", e)
            elif file_type == 'csv':
                file_name_ = case_name + '.csv'
                print("File name", file_name_)
                docx_file_path = os.path.join(csv_folder_path, file_name_)

                try:
                    file_name = Files(filename=file_name_, status='In Progress')
                    db.session.add(file_name)
                    db.session.commit()

                except Exception as e:
                    print("An error occurred while inserting data:", e)

            else:
                file_name_ = case_name + '.docx'
                docx_file_path = os.path.join(docx_folder_path, file_name_)
                print("File name", file_name_)

                try:
                    file_name = Files(filename=file_name_, status='In Progress')
                    db.session.add(file_name)
                    db.session.commit()
                except Exception as e:
                    print("An error occurred while inserting data:", e)

            case_names = case_name
            load_thread1 = threading.Thread(target=main_fun,
                                            args=(
                                                app, case_names, prompt_text, docx_file_path, file_name_,
                                                review_number))
            load_thread1.start()

        return render_template("Completed")

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# For Downloading the file
@app.route('/download/<filename>')
def download(filename):
    print(filename, "download")
    try:
        # filename = filename.replace("xlsx", 'docx') and filename.replace("xlsx", 'pdf')
        file_path = f'static/docx/{filename}'
        return send_file(file_path, as_attachment=True)
    except:
        try:
            file_path = f'static/csv/{filename}'
            return send_file(file_path, as_attachment=True)
        except:
            file_path = f'static/pdf/{filename}'
            return send_file(file_path, as_attachment=True)


# for deleting the file
@app.route('/delete', methods=['POST'])
def delete_file():
    try:
        file_id = request.form.get('file_id')
        file = Files.query.filter_by(id=file_id).first()
        db.session.delete(file)
        db.session.commit()
        return "File deleted successfully"
    except Exception as e:
        return f"An error occurred: {str(e)}", 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080, debug=True)
